from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth import authenticate, login, logout, get_user_model
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib.auth.forms import UserCreationForm
from django.contrib.auth.models import User
from django.contrib.auth.tokens import default_token_generator
from django.core.mail import send_mail
from django.urls import reverse
from django.utils.encoding import force_bytes, force_str
from django.utils.http import urlsafe_base64_encode, urlsafe_base64_decode
from django.template.loader import render_to_string
from django.contrib.sites.shortcuts import get_current_site
from .models import Suspect, ActivityLog
from .forms import SuspectRegistrationForm, CustomSignupForm
from .decorators import allowed_users
from .utils import log_activity
import weasyprint


# ------------------------------
# USER LOGIN (Non-admin only)
# ------------------------------
def user_login_view(request):
    error = None
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        user = authenticate(request, username=username, password=password)

        if user:
            if not user.is_superuser:
                login(request, user)
                ActivityLog.objects.create(user=user, action='User logged in')
                return redirect('home')
            else:
                error = 'Admins must use the Admin Login Page.'
        else:
            error = 'Invalid username or password.'

    return render(request, 'core/user_login.html', {'error': error})


# ------------------------------
# ADMIN LOGIN
# ------------------------------
def admin_login_view(request):
    error = None
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        user = authenticate(request, username=username, password=password)

        if user:
            if user.is_superuser:
                login(request, user)
                ActivityLog.objects.create(user=user, action='Admin logged in')
                return redirect('home')
            else:
                error = 'You do not have admin access.'
        else:
            error = 'Invalid username or password.'

    return render(request, 'core/admin_login.html', {'error': error})


# ------------------------------
# LOGOUT (All Users)
# ------------------------------
@login_required
def custom_logout(request):
    ActivityLog.objects.create(user=request.user, action='User logged out')
    logout(request)
    return redirect('user_login')

from .forms import AnonymousComplaintForm

# ------------------------------
# USER SIGNUP (With Email Activation)
# ------------------------------
def signup(request):
    if request.method == 'POST':
        form = CustomSignupForm(request.POST)
        if form.is_valid():
            user = form.save(commit=False)
            user.is_active = False  # User must activate via email
            user.save()

            current_site = get_current_site(request)
            subject = 'Activate Your Crime Tracking Account'
            token = default_token_generator.make_token(user)
            uid = urlsafe_base64_encode(force_bytes(user.pk))

            activation_link = reverse('activate', kwargs={'uidb64': uid, 'token': token})
            activation_url = f'http://{current_site.domain}{activation_link}'

            message = render_to_string('core/activation_email.html', {
                'user': user,
                'activation_url': activation_url,
            })

            send_mail(
                subject,
                message,
                'admin@crimetracking.com',
                [user.email],
                fail_silently=False,
            )

            messages.success(request, 'Account created! Please check your email to activate.')
            return redirect('user_login')
    else:
        form = CustomSignupForm()

    return render(request, 'core/signup.html', {'form': form})


# ------------------------------
# EMAIL ACCOUNT ACTIVATION
# ------------------------------
def activate(request, uidb64, token):
    try:
        uid = force_str(urlsafe_base64_decode(uidb64))
        user = User.objects.get(pk=uid)
    except (TypeError, ValueError, OverflowError, User.DoesNotExist):
        user = None

    if user and default_token_generator.check_token(user, token):
        user.is_active = True
        user.save()
        messages.success(request, 'Your account has been activated successfully!')
        return redirect('user_login')
    else:
        return render(request, 'core/activation_invalid.html')


# ------------------------------
# ANONYMOUS COMPLAINT FORM
# ------------------------------
from django.shortcuts import render, redirect
from .forms import AnonymousComplaintForm

def anonymous_complaint_view(request):
    if request.method == 'POST':
        form = AnonymousComplaintForm(request.POST, request.FILES)
        if form.is_valid():
            complaint = form.save(commit=False)
            complaint.is_anonymous = True  # Assign this since it's excluded from form
            complaint.save()
            return redirect('anonymous_thank_you')  # Make sure this URL exists
    else:
        form = AnonymousComplaintForm()
    return render(request, 'core/anonymous_complaint_form.html', {'form': form})


# ------------------------------
# THANK YOU PAGE (Anonymous Complaint)
# ------------------------------
def anonymous_thank_you(request):
    return render(request, 'core/anonymous_thank_you.html')


# ------------------------------
# HOME PAGE
# ------------------------------
def home(request):
    return render(request, 'core/home.html')

# ------------------------------
# MENU VIEWS
# ------------------------------
@login_required
def file_menu(request):
    return render(request, 'core/file_menu.html')

def report_menu(request):
    return render(request, 'core/report_menu.html')

def query_menu(request):
    return render(request, 'core/query_menu.html')

def help_menu(request):
    return render(request, 'core/help_menu.html')


# ------------------------------
# SUSPECT REGISTRATION (Admin only)
# ------------------------------
@login_required
@allowed_users(allowed_roles=['Admin'])
def suspect_registration(request):
    if request.method == 'POST':
        form = SuspectRegistrationForm(request.POST)
        if form.is_valid():
            form.save()
            ActivityLog.objects.create(user=request.user, action='Registered new suspect')
            messages.success(request, 'Suspect registered successfully.')
            return redirect('suspect_registration')
    else:
        form = SuspectRegistrationForm()

    return render(request, 'core/suspect_registration.html', {'form': form})


# ------------------------------
# SUSPECT SEARCH (Admin & Officer)
# ------------------------------
@login_required
@allowed_users(allowed_roles=['Admin', 'Officer'])
def suspect_search(request):
    query = request.GET.get('query', '')
    results = []

    if query:
        results = Suspect.objects.filter(
            full_name__icontains=query
        ) | Suspect.objects.filter(
            national_id__icontains=query
        )

    return render(request, 'core/suspect_search.html', {
        'query': query,
        'results': results
    })


# ------------------------------
# SUSPECT DETAIL / UPDATE
# ------------------------------
@login_required
@allowed_users(allowed_roles=['Admin', 'Officer'])
def suspect_detail(request, suspect_id):
    suspect = get_object_or_404(Suspect, id=suspect_id)

    if request.method == 'POST':
     form = SuspectRegistrationForm(request.POST, instance=suspect)
    if form.is_valid():
        suspect = form.save(commit=False)

        # Get or create the crime type dynamically
        crime_type_input = form.cleaned_data.get('crime_type', '').strip()
        crime_type_instance, created = CrimeCatalog.objects.get_or_create(
            crime_name__iexact=crime_type_input,
            defaults={'crime_name': crime_type_input}
        )

        # Assign the crime type
        suspect.crime_type = crime_type_instance
        suspect.save()

        messages.success(request, 'Suspect updated successfully.')
        return redirect('suspect_search')
    else:
        form = SuspectRegistrationForm(instance=suspect)

    return render(request, 'core/suspect_detail.html', {
        'form': form,
        'suspect': suspect
    })

# ------------------------------
# ADMIN CHECK (Utility)
# ------------------------------
def is_admin(user):
    return user.is_superuser

from .forms import CriminalForm, ComplaintForm, CrimeCatalogForm
from .models import CrimeCatalog, ActivityLog
from .utils import log_activity
from .decorators import allowed_users

# ------------------------------
# CRIMINAL REGISTRATION (Admin only)
# ------------------------------
from .models import Criminal, CrimeCatalog
from .forms import CriminalForm

@login_required
@allowed_users(allowed_roles=['Admin'])
def register_criminal(request):
    if not request.user.is_superuser:
        return render(request, 'core/register_criminal.html', {
            'form': None,
            'access_denied': True
        })

    if request.method == 'POST':
        form = CriminalForm(request.POST, request.FILES)
        if form.is_valid():
            criminal = form.save(commit=False)

            # Process the entered crime type string
            crime_type_name = form.cleaned_data['crime_type'].strip()

            # Get or create the CrimeCatalog instance
        crime_catalog = CrimeCatalog.objects.filter(crime_name__iexact=crime_type_name).first()
        if not crime_catalog:
            # Create a new crime type if not found
            crime_catalog = CrimeCatalog.objects.create(crime_name=crime_type_name)


            # Assign the CrimeCatalog instance to the criminal
            criminal.crime_type = crime_catalog
            criminal.save()

            ActivityLog.objects.create(user=request.user, action='Registered new criminal')
            messages.success(request, 'Criminal registered successfully.')
            return redirect('home')
    else:
        form = CriminalForm()

    return render(request, 'core/register_criminal.html', {
        'form': form,
        'access_denied': False
    })

# ------------------------------
# COMPLAINT FORM (Logged-in Users)
# ------------------------------
from .models import CrimeCatalog 

@login_required
def complaint_view(request):
    if request.method == 'POST':
        form = ComplaintForm(request.POST, request.FILES)
        if form.is_valid():
            complaint = form.save(commit=False)
            complaint.user = request.user

            # Get the string crime name submitted
            crime_name = form.cleaned_data['crime_type'].strip()

            # Create or get CrimeCatalog entry
            crime_catalog_entry = CrimeCatalog.objects.filter(
                crime_name__iexact=crime_name
            ).first()

            if not crime_catalog_entry:
                # Properly indented block
                crime_catalog_entry = CrimeCatalog.objects.create(
                    crime_name=crime_name,
                    description="Auto-added from user complaint.",
                    category='other',
                    severity='low'
                )

            # Assign the FK properly
            complaint.crime_type = crime_catalog_entry
            complaint.save()
            messages.success(request, 'Complaint submitted successfully!')
            return redirect('view_my_complaints')
        else:
            messages.error(request, 'Please correct the errors below.')
    else:
        form = ComplaintForm()

    return render(request, 'core/complaint_view.html', {
        'form': form,
        'crime_catalog_list': list(CrimeCatalog.objects.values_list('crime_name', flat=True))
    })

# ------------------------------
# CRIME CATALOG VIEWS (Admin only)
# ------------------------------
@login_required
@allowed_users(allowed_roles=['Admin'])
def crime_catalog(request):
    crimes = CrimeCatalog.objects.all()
    return render(request, 'core/crime_catalog.html', {'crimes': crimes})


@login_required
@allowed_users(allowed_roles=['Admin'])
def add_crime(request):
    if request.method == 'POST':
        form = CrimeCatalogForm(request.POST)
        if form.is_valid():
            form.save()
            ActivityLog.objects.create(user=request.user, action='Added new crime type')
            messages.success(request, 'New crime type added successfully.')
            return redirect('crime_catalog')
    else:
        form = CrimeCatalogForm()

    return render(request, 'core/add_crime.html', {'form': form})


@login_required
@allowed_users(allowed_roles=['Admin'])
def edit_crime(request, pk):
    crime = get_object_or_404(CrimeCatalog, pk=pk)
    if request.method == 'POST':
        form = CrimeCatalogForm(request.POST, instance=crime)
        if form.is_valid():
            form.save()
            ActivityLog.objects.create(user=request.user, action='Edited a crime type')
            messages.success(request, 'Crime type updated successfully.')
            return redirect('crime_catalog')
    else:
        form = CrimeCatalogForm(instance=crime)

    return render(request, 'core/add_crime.html', {'form': form})


@login_required
@allowed_users(allowed_roles=['Admin'])
def delete_crime(request, pk):
    crime = get_object_or_404(CrimeCatalog, pk=pk)
    if request.method == 'POST':
        crime.delete()
        messages.success(request, 'Crime type deleted.')
        return redirect('crime_catalog')

    return render(request, 'core/crime_catalog_confirm_delete.html', {'crime': crime})


# ------------------------------
# ADMIN-ONLY CHECK AND VIEW
# ------------------------------
def is_admin(user):
    return user.groups.filter(name='Admin').exists()

@login_required
@user_passes_test(is_admin)
def admin_only_view(request):
    # Placeholder view if needed later
    return render(request, 'core/admin_only.html')

from django.contrib.auth.models import User, Group
from .models import ActivityLog, Complaint, CrimeCatalog
from .decorators import allowed_users

# ------------------------------
# USER MANAGEMENT (Admin Only)
# ------------------------------
@login_required
@allowed_users(allowed_roles=['Admin'])
def manage_users(request):
    if not request.user.is_superuser:
        return redirect('home')  # Redirect non-superusers

    users = User.objects.all()
    groups = Group.objects.all()

    if request.method == 'POST':
        user_id = request.POST.get('user_id')
        group_id = request.POST.get('group_id')

        try:
            user = User.objects.get(id=user_id)
            group = Group.objects.get(id=group_id)
            user.groups.clear()
            user.groups.add(group)
        except (User.DoesNotExist, Group.DoesNotExist):
            pass  # Optionally handle/log this

        return redirect('manage_users')

    return render(request, 'core/manage_users.html', {
        'users': users,
        'groups': groups
    })


# ------------------------------
# GROUP CHECK UTILITY
# ------------------------------
def is_in_group(user, group_name):
    return user.groups.filter(name=group_name).exists()


# ------------------------------
# ACTIVITY LOG (All Users)
# ------------------------------
@login_required
def activity_log(request):
    logs = ActivityLog.objects.all().order_by('-timestamp')
    return render(request, 'core/activity_log.html', {'logs': logs})


# ------------------------------
# SYSTEM ACCESS PAGE
# ------------------------------
def system_access(request):
    return render(request, 'core/system_access.html')


# ------------------------------
# USER COMPLAINT VIEW
# ------------------------------
@login_required
def view_my_complaints(request):
    user_complaints = Complaint.objects.filter(user=request.user).order_by('-date_time_of_incident')
    return render(request, 'core/view_my_complaints.html', {'complaints': user_complaints})

# ------------------------------
# USER COMPLAINT DETAIL VIEW
# ------------------------------
@login_required
def complaint_user_detail(request, pk):
    complaint = get_object_or_404(Complaint, pk=pk, user=request.user)
    return render(request, 'core/complaint_user_detail.html', {'complaint': complaint})

# ------------------------------
# CRIME TYPES (Visible to Users)
# ------------------------------
@login_required
def crime_types_view(request):
    crimes = CrimeCatalog.objects.filter(is_active=True, reportable=True)

    category = request.GET.get('category')
    severity = request.GET.get('severity')

    if category:
        crimes = crimes.filter(category=category)
    if severity:
        crimes = crimes.filter(severity=severity)

    return render(request, 'core/crime_types_view.html', {
        'crime_types': crimes,
        'filter_category': category,
        'filter_severity': severity,
    })

from django.http import HttpResponse
from django.template.loader import get_template
from .models import Complaint
from xhtml2pdf import pisa
from docx import Document

# ------------------------------
# UTILITY: Admin Check
# ------------------------------
def is_admin(user):
    return user.is_superuser


# ------------------------------
# ALL COMPLAINTS VIEW (Admin Only)
# ------------------------------
from itertools import chain
from django.utils.timezone import localtime
from .models import Complaint, AnonymousComplaint

@login_required
@user_passes_test(is_admin)
def all_complaints_view(request):
    regular = Complaint.objects.all()
    anonymous = AnonymousComplaint.objects.all()

    # Tag the model type directly on each object
    for c in regular:
        c.model_type = 'regular'
        c.is_anonymous = False

    for a in anonymous:
        a.model_type = 'anonymous'
        a.is_anonymous = True

    # Combine and sort
    combined = sorted(
        chain(regular, anonymous),
        key=lambda x: x.date_time_of_incident if hasattr(x, 'date_time_of_incident') else x.incident_datetime,
        reverse=True
    )

    return render(request, 'core/all_complaints.html', {'complaints': combined})

# ------------------------------
# COMPLAINT DETAIL VIEW (Admin Only)
# ------------------------------
@login_required
@user_passes_test(is_admin)
def complaint_detail(request, pk):
    complaint = get_object_or_404(Complaint, pk=pk)
    return render(request, 'core/complaint_detail.html', {'complaint': complaint})

# ------------------------------
# ANONYMOUS COMPLAINT DETAIL VIEW (Admin Only)
# ------------------------------
@login_required
@user_passes_test(is_admin)
def anonymous_complaint_detail(request, pk):
    complaint = get_object_or_404(AnonymousComplaint, pk=pk)
    return render(request, 'core/anonymous_complaint_detail.html', {'complaint': complaint})


# ------------------------------
# USER COMPLAINT LIST VIEW (All Authenticated Users)
# ------------------------------
@login_required
def user_complaint_list_view(request):
    complaints = Complaint.objects.all()
    return render(request, 'core/user_complaint_list.html', {'complaints': complaints})


# ------------------------------
# EXPORT ALL COMPLAINTS - PDF
# ------------------------------
def export_complaints_pdf(request):
    complaints = Complaint.objects.all()
    template_path = 'core/export_complaints_pdf.html'
    context = {'complaints': complaints}

    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="complaints.pdf"'

    template = get_template(template_path)
    html = template.render(context)

    pisa_status = pisa.CreatePDF(html, dest=response)
    if pisa_status.err:
        return HttpResponse('PDF generation failed')
    return response


# ------------------------------
# EXPORT ALL COMPLAINTS - DOCX
# ------------------------------
def export_complaints_doc(request):
    complaints = Complaint.objects.all()
    document = Document()
    document.add_heading('All Complaints', 0)

    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Complainant'
    hdr_cells[1].text = 'Crime'
    hdr_cells[2].text = 'Status'
    hdr_cells[3].text = 'Emergency'

    for complaint in complaints:
        row = table.add_row().cells
        row[0].text = complaint.complainant_full_name
        row[1].text = str(complaint.crime_type)
        row[2].text = complaint.complaint_status
        row[3].text = 'Yes' if complaint.is_emergency else 'No'

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="complaints.docx"'
    document.save(response)
    return response

import openpyxl
from openpyxl.utils import get_column_letter
# import weasyprint

# ------------------------------
# EXPORT ALL COMPLAINTS - EXCEL (Admin)
# ------------------------------
def export_complaints_excel(request):
    complaints = Complaint.objects.all()
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Complaints"

    headers = ['Complainant', 'Crime', 'Status', 'Emergency']
    for col, header in enumerate(headers, 1):
        sheet[f"{get_column_letter(col)}1"] = header

    for row_num, complaint in enumerate(complaints, 2):
        sheet[f"A{row_num}"] = complaint.complainant_full_name
        sheet[f"B{row_num}"] = str(complaint.crime_type)
        sheet[f"C{row_num}"] = complaint.complaint_status
        sheet[f"D{row_num}"] = "Yes" if complaint.is_emergency else "No"

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=complaints.xlsx'
    workbook.save(response)
    return response


# ------------------------------
# COMPLAINT DETAIL (Generic View)
# ------------------------------
def complaint_detail_view(request, pk):
    complaint = get_object_or_404(Complaint, pk=pk)
    context = {'complaint': complaint}
    return render(request, 'report/complaint_detail.html', context)


# ------------------------------
# USER EXPORT – PDF
# ------------------------------
@login_required
def export_user_complaints_pdf(request):
    user_complaints = Complaint.objects.filter(user=request.user)
    template = get_template('core/user_complaints_pdf.html')
    html = template.render({'complaints': user_complaints})

    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="user_complaints.pdf"'

    weasyprint.HTML(string=html).write_pdf(response)
    return response


# ------------------------------
# USER EXPORT – DOCX
# ------------------------------
@login_required
def export_user_complaints_doc(request):
    user_complaints = Complaint.objects.filter(user=request.user)
    document = Document()
    document.add_heading('Your Complaint List', 0)

    for complaint in user_complaints:
        document.add_heading(f"Complaint #{complaint.id}", level=1)
        document.add_paragraph(f"Crime Type: {complaint.crime_type}")
        document.add_paragraph(f"Description: {complaint.crime_description}")
        document.add_paragraph(f"Status: {complaint.complaint_status}")
        document.add_paragraph(f"Emergency: {'Yes' if complaint.is_emergency else 'No'}")
        document.add_paragraph(f"Date: {complaint.incident_datetime.strftime('%Y-%m-%d %H:%M')}")

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="user_complaints.docx"'
    document.save(response)
    return response


# ------------------------------
# USER EXPORT – EXCEL
# ------------------------------
@login_required
def export_user_complaints_excel(request):
    user_complaints = Complaint.objects.filter(user=request.user)
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "User Complaints"

    headers = ['ID', 'Crime Type', 'Description', 'Status', 'Emergency', 'Date']
    sheet.append(headers)

    for complaint in user_complaints:
        sheet.append([
            complaint.id,
            str(complaint.crime_type),
            complaint.crime_description,
            complaint.complaint_status,
            'Yes' if complaint.is_emergency else 'No',
            complaint.incident_datetime.strftime('%Y-%m-%d %H:%M')
        ])

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="user_complaints.xlsx"'
    workbook.save(response)
    return response

# from weasyprint import HTML
# ------------------------------
# SINGLE COMPLAINT – EXPORT TO PDF
# ------------------------------
def export_single_complaint_pdf(request, pk):
    complaint = get_object_or_404(Complaint, pk=pk)

    # Ensure evidence has full URL for images or downloads
    if complaint.evidence:
        evidence_url = request.build_absolute_uri(complaint.evidence.url)
    else:
        evidence_url = None

    # Prepare a dictionary with all necessary complaint info
    complaint_data = {
        'full_name': complaint.full_name if not complaint.is_anonymous else "Anonymous",
        'email': complaint.email if not complaint.is_anonymous else "Hidden",
        'phone_number': complaint.phone_number if not complaint.is_anonymous else "Hidden",
        'crime_type': complaint.crime_type,
        'crime_description': complaint.crime_description,
        'location': complaint.location,
        'relationship': complaint.relationship,
        'suspect_description': complaint.suspect_description,
        'complaint_status': complaint.complaint_status,
        'is_anonymous': complaint.is_anonymous,
        'date_time_of_incident': complaint.date_time_of_incident,
        'evidence_url': evidence_url
    }

    # Use the same template you're using for HTML display
    template = get_template('core/user_complaints_pdf.html')
    html = template.render({'complaints': [complaint_data]})  # render as a list

    # Generate PDF
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename=complaint_{pk}.pdf'
    weasyprint.HTML(string=html, base_url=request.build_absolute_uri()).write_pdf(response)

    return response


# ------------------------------
# SINGLE COMPLAINT – EXPORT TO DOCX
# ------------------------------
from docx.shared import Inches
import os

def export_single_complaint_doc(request, pk):
    complaint = get_object_or_404(Complaint, pk=pk)
    doc = Document()
    doc.add_heading('Complaint Report', 0)

    doc.add_paragraph(f"Complainant: {complaint.full_name}")
    doc.add_paragraph(f"Email: {complaint.email}")
    doc.add_paragraph(f"Phone: {complaint.phone_number}")
    doc.add_paragraph(f"Crime Type: {complaint.crime_type}")
    doc.add_paragraph(f"Location: {complaint.location}")
    doc.add_paragraph(f"Anonymous: {'Yes' if complaint.is_anonymous else 'No'}")
    doc.add_paragraph(f"Status: {complaint.complaint_status}")
    doc.add_paragraph(f"Description: {complaint.crime_description}")
    doc.add_paragraph(f"Relationship to Suspect: {complaint.relationship}")
    doc.add_paragraph(f"Suspect Description: {complaint.suspect_description or 'None'}")

    if complaint.evidence:
        if complaint.evidence.name.lower().endswith(('.jpg', '.jpeg', '.png')):
            evidence_path = complaint.evidence.path
            if os.path.exists(evidence_path):
                doc.add_paragraph("Evidence Image:")
                doc.add_picture(evidence_path, width=Inches(4))
            else:
                doc.add_paragraph("Evidence image file not found.")
        else:
            doc.add_paragraph(f"Evidence File: {complaint.evidence.url}")
    else:
        doc.add_paragraph("Evidence: None")

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=complaint_{pk}.docx'
    doc.save(response)
    return response


# ------------------------------
# SINGLE COMPLAINT – EXPORT TO EXCEL
# ------------------------------
def export_single_complaint_excel(request, pk):
    complaint = get_object_or_404(Complaint, pk=pk)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Complaint Report"

    ws.append(["Field", "Value"])
    ws.append(["Complainant", complaint.full_name])
    ws.append(["Email", complaint.email])
    ws.append(["Phone", complaint.phone_number])
    ws.append(["Crime Type", str(complaint.crime_type)])
    ws.append(["Location", complaint.location])
    ws.append(["Anonymous", 'Yes' if complaint.is_anonymous else 'No'])
    ws.append(["Status", complaint.complaint_status])
    ws.append(["Relationship to Suspect", complaint.relationship])
    ws.append(["Date/Time of Incident", complaint.date_time_of_incident.strftime('%Y-%m-%d %H:%M')])
    ws.append(["Description of Crime", complaint.crime_description])
    ws.append(["Suspect Description", complaint.suspect_description or 'N/A'])
    ws.append(["Evidence File", complaint.evidence.url if complaint.evidence else 'None'])

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename=complaint_{pk}.xlsx'
    wb.save(response)
    return response
# ------------------------------
# SUSPECT LIST AND DETAIL VIEWS (Admin Only)
# ------------------------------

@allowed_users(allowed_roles=['Admin'])
@login_required
def suspect_list(request):
    suspects = Suspect.objects.all().order_by('date_of_birth')  # Optional sorting
    return render(request, 'core/suspect_list.html', {'suspects': suspects})

@login_required
@allowed_users(allowed_roles=['Admin', 'Officer'])
def suspect_detail_list(request, pk):
    suspect = get_object_or_404(Suspect, pk=pk)

    if request.method == 'POST':
        form = SuspectRegistrationForm(request.POST, request.FILES, instance=suspect)
        if form.is_valid():
            suspect_instance = form.save(commit=False)
            crime_type_input = form.cleaned_data['crime_type']

            # Ensure crime type is either fetched or created
            crime_obj, created = CrimeCatalog.objects.get_or_create(
                crime_name__iexact=crime_type_input,
                defaults={'crime_name': crime_type_input, 'description': '', 'category': 'other'}
            )
            suspect_instance.crime_type = crime_obj
            suspect_instance.save()
            messages.success(request, 'Suspect updated successfully.')
            return redirect('suspect_search')
    else:
        form = SuspectRegistrationForm(instance=suspect)

    return render(request, 'core/suspect_detail_list.html', {
        'form': form,
        'suspect': suspect
    })


import io
from django.http import FileResponse, HttpResponse
from django.contrib.auth.decorators import login_required
from xhtml2pdf import pisa
import openpyxl
from docx import Document
from django.template.loader import get_template
from .models import Suspect
from .decorators import allowed_users  # Assuming you have this decorator
import csv

# PDF Export
import os
from django.conf import settings

@login_required
@allowed_users(allowed_roles=['Admin'])
def export_suspects_pdf(request):
    suspects = Suspect.objects.all()
    template_path = 'core/suspect_pdf.html'
    context = {
        'suspects': suspects,
        'media_path': settings.MEDIA_ROOT  # Needed for resolving image paths
    }

    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="suspects.pdf"'

    template = get_template(template_path)
    html = template.render(context)

    pisa_status = pisa.CreatePDF(
        io.BytesIO(html.encode("utf-8")),
        dest=response,
        encoding='utf-8'
    )

    if pisa_status.err:
        return HttpResponse('Error generating PDF', status=500)
    return response


# Excel Export
@login_required
@allowed_users(allowed_roles=['Admin'])
def export_suspects_excel(request):
    suspects = Suspect.objects.all()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Suspects"

    ws.append([
        'Full Name', 'Nickname', 'Date of Birth', 'Gender', 'Address',
        'Phone Number', 'Email', 'National ID', 'Crime Type',
        'Known Associates', 'Previous Criminal Record', 'Current Status'
    ])

    for s in suspects:
        ws.append([
            s.full_name, s.nickname, s.date_of_birth.strftime('%Y-%m-%d') if s.date_of_birth else '',
            s.gender, s.address, s.phone_number, s.email, s.national_id,
            str(s.crime_type), s.known_associates, s.previous_criminal_record, s.current_status
        ])

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="suspects.xlsx"'
    wb.save(response)
    return response

# Word Export
import os
from django.conf import settings
from docx.shared import Inches

@login_required
@allowed_users(allowed_roles=['Admin'])
def export_suspects_word(request):
    suspects = Suspect.objects.all()
    doc = Document()
    doc.add_heading('Suspect List', 0)

    for s in suspects:
        doc.add_paragraph(f"""
Full Name: {s.full_name}
Nickname: {s.nickname}
Date of Birth: {s.date_of_birth.strftime('%Y-%m-%d') if s.date_of_birth else ''}
Gender: {s.gender}
Address: {s.address}
Phone Number: {s.phone_number}
Email: {s.email}
National ID: {s.national_id}
Crime Type: {s.crime_type}
Known Associates: {s.known_associates}
Previous Criminal Record: {s.previous_criminal_record}
Current Status: {s.current_status}
        """)
        if s.picture and os.path.isfile(s.picture.path):
            doc.add_picture(s.picture.path, width=Inches(1.5))
        else:
            doc.add_paragraph("No image available.")
        doc.add_paragraph("------------------------------------------------------------")

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="suspects.docx"'
    doc.save(response)
    return response

from .models import Criminal

def criminal_list(request):
    criminals = Criminal.objects.all()
    return render(request, 'core/criminal_list.html', {'criminals': criminals})

def criminal_detail(request, pk):
    criminal = get_object_or_404(Criminal, pk=pk)
    return render(request, 'core/criminal_detail.html', {'criminal': criminal})

# Export to PDF
def export_criminals_pdf(request):
    # Placeholder logic — implement real PDF generation (e.g., using ReportLab or WeasyPrint)
    return HttpResponse("PDF export placeholder", content_type="application/pdf")

# Export to Excel
def export_criminals_excel(request):
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="criminals.csv"'

    writer = csv.writer(response)
    writer.writerow(['Full Name', 'Criminal ID', 'DOB', 'Gender', 'Address', 'Phone', 'Email', 'Date of Conviction', 'Crime Type', 'Sentence', 'Associates'])

    for c in Criminal.objects.all():
        writer.writerow([c.full_name, c.criminal_id_number, c.date_of_birth, c.gender, c.address, c.phone_number, c.email, c.date_of_conviction, c.crime_type, c.sentence_details, c.known_associates])

    return response

# Export to Word
def export_criminals_word(request):
    response = HttpResponse(content_type='application/msword')
    response['Content-Disposition'] = 'attachment; filename="criminals.doc"'
    response.write("Criminal List Export Placeholder")  # Replace with real Word export logic
    return response

from django.http import JsonResponse
from .models import CrimeCatalog
from django.views.decorators.http import require_GET

@require_GET
def crime_autocomplete(request):
    term = request.GET.get('term', '')
    crimes = CrimeCatalog.objects.filter(crime_name__icontains=term)[:10]
    results = list(crimes.values_list('crime_name', flat=True))
    return JsonResponse(results, safe=False)

from django.shortcuts import render

def test_template(request):
    return render(request, "y")